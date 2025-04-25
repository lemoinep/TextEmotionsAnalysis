import os
import glob
import argparse
import pandas as pd
import matplotlib.pyplot as plt
from nrclex import NRCLex
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

import win32com.client  # pywin32 for reading .doc files

def read_file_content(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif ext == ".docx":
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)
    elif ext == ".doc":
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path)
            text = doc.Content.Text
            doc.Close(False)
            word.Quit()
            return text
        except Exception as e:
            print(f"Error reading .doc file {file_path}: {e}")
            return ""
    else:
        print(f"Unsupported file extension: {ext} for file {file_path}")
        return ""

def get_coeff_emotion(emotions_list):
    coeff = 0
    emotions = [e[0].lower() for e in emotions_list]
    emotion_weights = {
        "fear": -1,
        "anger": -1,
        "anticip": 1,
        "trust": 1,
        "surprise": 1,
        "positive": 1,
        "negative": -1,
        "sadness": -1,
        "disgust": -1,
        "joy": 1
    }
    for emotion in emotions:
        if emotion in emotion_weights:
            coeff += emotion_weights[emotion]
    return coeff

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--Path", type=str, default=None, help="Path to the folder containing files")
    parser.add_argument("--QRebuild", type=int, default=0, help="QRebuild option (not used here)")
    args = parser.parse_args()

    QView = False

    current_directory = os.getcwd()
    path_dataset = args.Path if args.Path else os.path.join(current_directory, "DATA")

    if QView:
        print(f"Dataset path = {path_dataset}")

    # Collect .txt, .doc, .docx files
    list_files = []
    for ext in ("*.txt", "*.doc", "*.docx"):
        list_files.extend(glob.glob(os.path.join(path_dataset, ext)))

    color_map_emotion = {
        "fear": "red",
        "anger": "darkred",
        "anticip": "gray",
        "trust": "turquoise",
        "surprise": "pink",
        "positive": "limegreen",
        "negative": "red",
        "sadness": "violet",
        "disgust": "goldenrod",
        "joy": "yellow"
    }

    color_map_score = [
        (lambda s: s <= -4, WD_COLOR_INDEX.DARK_RED),
        (lambda s: s <= -3, WD_COLOR_INDEX.RED),
        (lambda s: s <= -1, WD_COLOR_INDEX.VIOLET),
        (lambda s: s == 0, WD_COLOR_INDEX.WHITE),
        (lambda s: s >= 4, WD_COLOR_INDEX.BRIGHT_GREEN),
        (lambda s: s >= 3, WD_COLOR_INDEX.GREEN),
        (lambda s: s >= 1, WD_COLOR_INDEX.GRAY_25)
    ]

    for file_path in list_files:
        if QView:
            print(f"\nProcessing file: {file_path}")

        text_to_read = read_file_content(file_path)
        if not text_to_read.strip():
            print(f"Warning: file {file_path} is empty or could not be read.")
            continue

        if QView:
            print("\nRead text:\n", text_to_read, "\n")

        text_object = NRCLex(text_to_read)

        file_base_name, _ = os.path.splitext(file_path)

        # CSV report affect_dict
        affect_dict = text_object.affect_dict
        df_affect = pd.DataFrame(list(affect_dict.items()), columns=["Emotion", "Count"])
        csv_analysis_path = f"{file_base_name}_Report_Analysis.csv"
        df_affect.to_csv(csv_analysis_path, index=False)

        # CSV report affect_frequencies
        affect_freq = text_object.affect_frequencies
        df_freq = pd.DataFrame(list(affect_freq.items()), columns=["Emotion", "Frequency"])
        csv_freq_path = f"{file_base_name}_Report.csv"
        df_freq.to_csv(csv_freq_path, index=False)

        # Bar plot
        X = df_freq["Emotion"].tolist()
        Y = df_freq["Frequency"].tolist()

        plt.figure(figsize=(10, 6))
        plt.ylim(min(Y) * 0.95, max(Y) * 1.05)
        bars = plt.bar(X, Y)

        for bar, emotion in zip(bars, X):
            bar.set_color(color_map_emotion.get(emotion.lower(), "gray"))

        plt.title("Text Emotion Analysis")
        plt.ylabel("Frequency")
        plt.xticks(rotation=45)
        plt.tight_layout()

        img_path = f"{file_base_name}_Report.jpg"
        plt.savefig(img_path)
        plt.close()

        # Word document with sentence highlighting
        doc = docx.Document()
        doc.add_heading("Text Emotion Analysis", level=0)

        styles = doc.styles
        char_style = styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
        font = char_style.font
        font.size = Pt(10)
        font.name = 'Times New Roman'

        for sentence in text_object.sentences:
            phrase = str(sentence) + " "
            phrase_obj = NRCLex(phrase)
            top_emotions = phrase_obj.top_emotions

            score = get_coeff_emotion(top_emotions)

            highlight_color = WD_COLOR_INDEX.WHITE
            for cond, color in color_map_score:
                if cond(score):
                    highlight_color = color
                    break

            para = doc.add_paragraph()
            run = para.add_run(phrase, style='CommentsStyle')
            run.font.highlight_color = highlight_color

        doc_path = f"{file_base_name}_Report_Colors.docx"
        doc.save(doc_path)

if __name__ == "__main__":
    main()
